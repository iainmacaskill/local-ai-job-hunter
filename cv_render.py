"""Render an ATS-safe "screening" CV: plain, single-column, keyword-aligned.

The screening CV has one job — pass automated keyword/semantic screening — so it
is deliberately *unstyled*: single column, standard headings, real text, no
tables or graphics (which trip ATS parsers). The keyword-mirrored content is
produced by the drafter (``draft_cv``), which reads the job description and maps
its terms to the candidate's genuine experience; this module renders that into a
parser-friendly .docx and reports keyword coverage.

Honesty rule: only include keywords for skills the candidate actually has. Use
``keyword_coverage()`` to surface real gaps rather than inventing terms — fake
keywords get caught and collapse at interview.

Internalised from jobtracker's ``screening_cv`` (A2): it now takes the profile
dict directly and has no dependency on jobtracker's ``config``/``cv_builder``.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from cv_profile import OUTPUT_DIR

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Calibri"


def _run(p, text, size=11, bold=False):
    r = p.add_run(str(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = BLACK
    return r


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _run(p, text.upper(), size=11, bold=True)


def _safe(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "", (name or "")).strip() or "Role"


def cv_fulltext(screening: dict, profile: dict) -> str:
    """Flatten everything a parser would read, for a coverage check."""
    parts = [screening.get("target_title", ""), screening.get("summary", "")]
    parts += list(screening.get("core_skills", []))
    for job in screening.get("experience") or profile.get("jobs", []):
        parts += [job.get("title", ""), job.get("company", "")]
        parts += list(job.get("bullets", []))
    parts += [profile.get("certifications", ""), profile.get("education", "")]
    return "\n".join(str(p) for p in parts)


def keyword_coverage(keywords, text: str) -> dict:
    """Report which JD keywords literally appear in ``text`` (covered) or not.

    Short/acronym terms match on word boundaries (so "AI" doesn't match "detail");
    longer phrases match as substrings.
    """
    low = text.lower()
    covered, missing = [], []
    for kw in keywords:
        k = str(kw).strip()
        if not k:
            continue
        kl = k.lower()
        if len(k) <= 4 or k.isupper():  # acronym/short term — match whole word only
            found = re.search(rf"\b{re.escape(kl)}\b", low) is not None
        else:  # phrase — plain substring match
            found = kl in low
        (covered if found else missing).append(k)
    total = len(covered) + len(missing)
    pct = round(100 * len(covered) / max(total, 1))
    return {"covered": covered, "missing": missing, "pct": pct}


def generate_screening_cv(role: dict, screening: dict, profile: dict, out_dir=None) -> Path:
    """Render an ATS-safe screening CV from a drafter-produced ``screening`` payload.

    ``screening`` carries the JD-mirrored content:
        target_title : str        — mirrors the advert's job title
        summary      : str        — keyword-rich professional summary
        core_skills  : list[str]  — the keyword block (rendered as plain text)
        experience   : list[dict] — optional; {title, company, dates, bullets}
                                     falls back to the profile's jobs if omitted
    Structural data (name, contact, certs, education) comes from the profile.
    """
    out_dir = Path(out_dir) if out_dir else OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    target = (screening.get("target_title") or role.get("title") or "Role").strip()
    path = out_dir / f"{profile.get('name', 'CV')} - Screening - {_safe(target)}.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)

    # Header — plain text lines, no columns or tables.
    _run(doc.add_paragraph(), profile.get("name", ""), size=16, bold=True)
    _run(doc.add_paragraph(), target, size=12)
    _run(doc.add_paragraph(), profile.get("contact", ""), size=10)

    _heading(doc, "Professional Summary")
    _run(doc.add_paragraph(), screening.get("summary", ""))

    _heading(doc, "Core Skills")
    skills = [s for s in screening.get("core_skills", []) if str(s).strip()]
    if skills:
        _run(doc.add_paragraph(), ", ".join(skills))

    _heading(doc, "Professional Experience")
    for job in screening.get("experience") or profile.get("jobs", []):
        _run(doc.add_paragraph(), f"{job.get('title', '')}, {job.get('company', '')}", bold=True)
        if job.get("dates"):
            _run(doc.add_paragraph(), job["dates"], size=10)
        for bullet in job.get("bullets", []):
            _run(doc.add_paragraph(style="List Bullet"), bullet)

    _heading(doc, "Education & Certifications")
    if profile.get("certifications"):
        _run(doc.add_paragraph(), profile["certifications"])
    if profile.get("education"):
        _run(doc.add_paragraph(), profile["education"])

    # Honest document metadata. python-docx's bundled template otherwise leaves
    # author='python-docx' and a 2013 created date — a machine-generation
    # fingerprint visible to anyone who opens File > Properties.
    now = datetime.now(timezone.utc)
    cp = doc.core_properties
    cp.author = profile.get("name", "")
    cp.last_modified_by = profile.get("name", "")
    cp.created = now
    cp.modified = now
    cp.title = f"{profile.get('name', 'CV')} - {target}"

    doc.save(str(path))
    return path
