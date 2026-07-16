"""S4: draft a tailored cover letter from a JD using the local model.

Same principles as the CV drafter: the local model writes the letter body using
ONLY genuine profile facts, in house style (British English, no em dashes), and
honesty.verify_text flags any figure not in the profile or any em dash. Renders a
.docx (for attaching) and a .txt (for pasting into an email or portal). The user
reviews and sends -- nothing is sent from here.
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
import sys
from datetime import timezone
from pathlib import Path

import honesty
import settings
from local_llm import LocalLLM

settings.wire_jobtracker()
import screening_cv  # noqa: E402  (reused for load_profile)
from docx import Document  # noqa: E402

OUTPUT_DIR = Path(__file__).resolve().parent / "outputs"

STYLE = (
    "Write in British English. Do not use em dashes (use commas, colons or "
    "parentheses instead). Do not invent employers, metrics or claims that are "
    "not in the facts provided. Warm and professional, never effusive."
)


def _safe(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "", (name or "")).strip() or "Untitled"


def _stem(company: str | None, role_title: str | None) -> str:
    return f"Cover Letter - {_safe(company or 'Company')} - {_safe(role_title or 'Role')}"


def _facts(profile: dict) -> str:
    roles = "; ".join(
        f"{j.get('title')} at {j.get('company')} ({j.get('dates')})"
        for j in profile.get("jobs", [])
    )
    achievements = "; ".join(profile.get("achievements", []))
    comps = ", ".join(profile.get("competencies", [])[:12])
    return (
        f"Roles: {roles}\n"
        f"Key achievements: {achievements}\n"
        f"Competencies: {comps}\n"
        f"Certifications: {profile.get('certifications', '')}"
    )


def draft_paragraphs(
    jd_text: str, role_title: str | None, company: str | None, profile: dict, llm
) -> list[str]:
    out = llm.complete_json(
        system=(
            "Write the body of a concise cover letter as exactly three short paragraphs, "
            "using ONLY the candidate facts provided. Paragraph one: the role applied for "
            "and the headline reason the candidate fits. Paragraph two: two or three "
            "specific, genuine achievements relevant to this advert. Paragraph three: "
            "enthusiasm and a brief close. "
            'Return JSON: {"paragraphs": array of exactly 3 strings}. ' + STYLE
        ),
        user=(
            f"JOB ADVERT:\n{jd_text}\n\n"
            f"ROLE: {role_title or '(use the advert job title)'}\n"
            f"ORGANISATION: {company or '(use the advert; otherwise say your organisation)'}\n\n"
            f"CANDIDATE FACTS:\n{_facts(profile)}"
        ),
        max_tokens=800,
    )
    paras = [str(p).strip() for p in out.get("paragraphs", []) if str(p).strip()]
    return paras[:3]


def _stamp(doc: Document, name: str, title: str) -> None:
    now = dt.datetime.now(timezone.utc)
    cp = doc.core_properties
    cp.author = cp.last_modified_by = name
    cp.created = cp.modified = now
    cp.title = title


def render_docx(
    profile: dict, company: str | None, role_title: str | None,
    paragraphs: list[str], out_dir: Path,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    name = profile.get("name", "")
    path = out_dir / f"{_stem(company, role_title)}.docx"

    doc = Document()
    doc.add_paragraph().add_run(name).bold = True
    doc.add_paragraph(profile.get("contact", ""))
    doc.add_paragraph("")
    doc.add_paragraph(dt.date.today().strftime("%d %B %Y"))
    doc.add_paragraph("")
    doc.add_paragraph("Dear Hiring Manager,")
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.add_paragraph("")
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph(name)

    _stamp(doc, name, f"Cover letter - {role_title or ''}".strip())
    doc.save(str(path))
    return path


def render_txt(
    profile: dict, company: str | None, role_title: str | None,
    paragraphs: list[str], out_dir: Path,
) -> Path:
    name = profile.get("name", "")
    path = out_dir / f"{_stem(company, role_title)}.txt"
    lines = [
        name, profile.get("contact", ""), "",
        dt.date.today().strftime("%d %B %Y"), "",
        "Dear Hiring Manager,", "",
        *[p + "\n" for p in paragraphs],
        "Yours sincerely,", name,
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def draft_cover_letter(
    jd_text: str,
    role_title: str | None = None,
    company: str | None = None,
    profile: dict | None = None,
    llm=None,
    out_dir: Path | None = None,
) -> dict:
    profile = profile or screening_cv.load_profile()
    llm = llm or LocalLLM(base_url=settings.LLM_BASE_URL, model=settings.LLM_MODEL)
    out_dir = Path(out_dir) if out_dir else OUTPUT_DIR

    paragraphs = draft_paragraphs(jd_text, role_title, company, profile, llm)
    report = honesty.verify_text("\n".join(paragraphs), profile, what="cover letter")
    docx = render_docx(profile, company, role_title, paragraphs, out_dir)
    txt = render_txt(profile, company, role_title, paragraphs, out_dir)
    return {"docx": docx, "txt": txt, "paragraphs": paragraphs, "honesty": report}


def _main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Draft a cover letter from a JD (local model).")
    ap.add_argument("--jd-file", required=True, help="path to a JD text file, or - for stdin")
    ap.add_argument("--title", default=None, help="target role title")
    ap.add_argument("--company", default=None, help="organisation name")
    args = ap.parse_args(argv)
    jd = sys.stdin.read() if args.jd_file == "-" else Path(args.jd_file).read_text(encoding="utf-8")
    res = draft_cover_letter(jd, role_title=args.title, company=args.company)
    print(f"saved:   {res['docx']}")
    print(f"         {res['txt']}")
    print(f"honesty: {res['honesty'].summary()}")
    for warn in res["honesty"].warnings:
        print(f"  review {warn}")
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
