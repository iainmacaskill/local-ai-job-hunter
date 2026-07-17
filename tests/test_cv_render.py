"""Cover the internalised CV engine (A2): ATS structure, honest metadata, coverage."""

from docx import Document

import cv_render
from cv_profile import load_profile

SCREENING = {
    "target_title": "AI Delivery Manager",
    "summary": "Delivery leader with AI and data programme experience.",
    "core_skills": ["AI", "data delivery", "Agile", "PRINCE2"],
    "experience": [
        {"title": "Programme Manager", "company": "Acme", "dates": "2020-2024",
         "bullets": ["Led AI delivery.", "Ran Agile teams."]},
    ],
}


def test_screening_cv_is_ats_safe_and_honestly_stamped(tmp_path):
    prof = load_profile()
    path = cv_render.generate_screening_cv(
        {"title": "AI Delivery Manager"}, SCREENING, prof, out_dir=tmp_path
    )
    assert path.exists() and path.suffix == ".docx"

    doc = Document(str(path))
    # ATS parsers choke on tables/columns — the screening CV must have none.
    assert not doc.tables
    # Honest document metadata: author is the candidate, not python-docx's default.
    assert doc.core_properties.author == prof["name"]
    assert doc.core_properties.last_modified_by == prof["name"]
    # The keyword-mirrored summary made it into the rendered text.
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI and data programme experience" in text


def test_keyword_coverage_word_boundary_vs_substring():
    # Acronyms/short terms match whole words only ("AI" must not match "detail").
    assert cv_render.keyword_coverage(["AI"], "attention to detail")["pct"] == 0
    assert cv_render.keyword_coverage(["AI"], "AI delivery")["pct"] == 100
    # Longer phrases match as substrings.
    cov = cv_render.keyword_coverage(["stakeholder management"], "strong stakeholder management")
    assert cov["covered"] == ["stakeholder management"] and cov["missing"] == []


def test_cv_fulltext_flattens_skills_and_bullets():
    prof = load_profile()
    text = cv_render.cv_fulltext(SCREENING, prof)
    assert "PRINCE2" in text and "Led AI delivery." in text
